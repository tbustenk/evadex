"""Integration tests for evadex generate CLI command."""
from __future__ import annotations

import ast
import csv
import json
from pathlib import Path

import pytest
from click.testing import CliRunner

from evadex.cli.app import main


# ── Helpers ────────────────────────────────────────────────────────────────────

def _invoke(runner: CliRunner, args: list[str]):
    return runner.invoke(main, ["generate"] + args)


def _base_args(fmt: str, output: str, count: int = 5) -> list[str]:
    return [
        "--format", fmt,
        "--category", "credit_card",
        "--count", str(count),
        "--seed", "42",
        "--output", output,
    ]


# ── Exit codes and basic sanity ────────────────────────────────────────────────

def test_generate_csv_exits_zero(tmp_path):
    runner = CliRunner()
    out = str(tmp_path / "out.csv")
    result = _invoke(runner, _base_args("csv", out))
    assert result.exit_code == 0, result.output


def test_generate_txt_exits_zero(tmp_path):
    runner = CliRunner()
    out = str(tmp_path / "out.txt")
    result = _invoke(runner, _base_args("txt", out))
    assert result.exit_code == 0, result.output


def test_generate_xlsx_exits_zero(tmp_path):
    runner = CliRunner()
    out = str(tmp_path / "out.xlsx")
    result = _invoke(runner, _base_args("xlsx", out))
    assert result.exit_code == 0, result.output


def test_generate_docx_exits_zero(tmp_path):
    runner = CliRunner()
    out = str(tmp_path / "out.docx")
    result = _invoke(runner, _base_args("docx", out))
    assert result.exit_code == 0, result.output


def test_generate_pdf_exits_zero(tmp_path):
    runner = CliRunner()
    out = str(tmp_path / "out.pdf")
    result = _invoke(runner, _base_args("pdf", out))
    assert result.exit_code == 0, result.output


# ── Files are created ──────────────────────────────────────────────────────────

def test_csv_file_created(tmp_path):
    runner = CliRunner()
    out = tmp_path / "out.csv"
    _invoke(runner, _base_args("csv", str(out)))
    assert out.exists()


def test_txt_file_created(tmp_path):
    runner = CliRunner()
    out = tmp_path / "out.txt"
    _invoke(runner, _base_args("txt", str(out)))
    assert out.exists()


def test_xlsx_file_created(tmp_path):
    runner = CliRunner()
    out = tmp_path / "out.xlsx"
    _invoke(runner, _base_args("xlsx", str(out)))
    assert out.exists()


def test_docx_file_created(tmp_path):
    runner = CliRunner()
    out = tmp_path / "out.docx"
    _invoke(runner, _base_args("docx", str(out)))
    assert out.exists()


def test_pdf_file_created(tmp_path):
    runner = CliRunner()
    out = tmp_path / "out.pdf"
    _invoke(runner, _base_args("pdf", str(out)))
    assert out.exists()


# ── CSV content ────────────────────────────────────────────────────────────────

def test_csv_has_correct_headers(tmp_path):
    runner = CliRunner()
    out = tmp_path / "out.csv"
    _invoke(runner, _base_args("csv", str(out), count=3))
    with open(out, encoding="utf-8") as fh:
        reader = csv.DictReader(fh)
        assert "category" in reader.fieldnames
        assert "plain_value" in reader.fieldnames
        assert "variant_value" in reader.fieldnames
        assert "technique" in reader.fieldnames
        assert "embedded_text" in reader.fieldnames


def test_csv_count_respected(tmp_path):
    runner = CliRunner()
    out = tmp_path / "out.csv"
    _invoke(runner, _base_args("csv", str(out), count=7))
    with open(out, encoding="utf-8") as fh:
        rows = list(csv.DictReader(fh))
    assert len(rows) == 7


def test_csv_category_column_matches(tmp_path):
    runner = CliRunner()
    out = tmp_path / "out.csv"
    _invoke(runner, _base_args("csv", str(out), count=5))
    with open(out, encoding="utf-8") as fh:
        rows = list(csv.DictReader(fh))
    assert all(r["category"] == "credit_card" for r in rows)


# ── TXT content ───────────────────────────────────────────────────────────────

def test_txt_contains_dlp_header(tmp_path):
    runner = CliRunner()
    out = tmp_path / "out.txt"
    _invoke(runner, _base_args("txt", str(out)))
    content = out.read_text(encoding="utf-8")
    assert "DLP TEST DOCUMENT" in content


def test_txt_contains_category_section(tmp_path):
    runner = CliRunner()
    out = tmp_path / "out.txt"
    _invoke(runner, _base_args("txt", str(out)))
    content = out.read_text(encoding="utf-8")
    assert "Credit Card" in content


# ── XLSX structure ────────────────────────────────────────────────────────────

def test_xlsx_has_summary_sheet(tmp_path):
    import openpyxl
    runner = CliRunner()
    out = tmp_path / "out.xlsx"
    _invoke(runner, _base_args("xlsx", str(out)))
    wb = openpyxl.load_workbook(out)
    assert "Summary" in wb.sheetnames


def test_xlsx_has_category_sheet(tmp_path):
    import openpyxl
    runner = CliRunner()
    out = tmp_path / "out.xlsx"
    _invoke(runner, _base_args("xlsx", str(out)))
    wb = openpyxl.load_workbook(out)
    # At least two sheets: Summary + credit_cards
    assert len(wb.sheetnames) >= 2


def test_xlsx_row_count_matches(tmp_path):
    import openpyxl
    runner = CliRunner()
    out = tmp_path / "out.xlsx"
    _invoke(runner, _base_args("xlsx", str(out), count=8))
    wb = openpyxl.load_workbook(out)
    # Find the credit_card sheet (not Summary)
    data_sheet = next(ws for ws in wb.worksheets if ws.title != "Summary")
    # Row 1 = header, rows 2..9 = data
    assert data_sheet.max_row == 9  # 1 header + 8 data rows


# ── DOCX structure ────────────────────────────────────────────────────────────

def test_docx_is_valid_docx(tmp_path):
    from docx import Document
    runner = CliRunner()
    out = tmp_path / "out.docx"
    _invoke(runner, _base_args("docx", str(out)))
    doc = Document(str(out))
    assert len(doc.paragraphs) > 0


def test_docx_contains_title(tmp_path):
    from docx import Document
    runner = CliRunner()
    out = tmp_path / "out.docx"
    _invoke(runner, _base_args("docx", str(out)))
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "DLP Test Document" in text


# ── PDF structure ─────────────────────────────────────────────────────────────

def test_pdf_starts_with_pdf_magic(tmp_path):
    runner = CliRunner()
    out = tmp_path / "out.pdf"
    _invoke(runner, _base_args("pdf", str(out)))
    magic = out.read_bytes()[:4]
    assert magic == b"%PDF"


# ── --seed reproducibility ─────────────────────────────────────────────────────

def test_seed_produces_identical_csv(tmp_path):
    runner = CliRunner()
    a = tmp_path / "a.csv"
    b = tmp_path / "b.csv"
    _invoke(runner, _base_args("csv", str(a), count=10))
    _invoke(runner, _base_args("csv", str(b), count=10))
    assert a.read_text(encoding="utf-8") == b.read_text(encoding="utf-8")


def test_different_seeds_produce_different_csv(tmp_path):
    runner = CliRunner()
    a = tmp_path / "a.csv"
    b = tmp_path / "b.csv"
    args_a = _base_args("csv", str(a), count=10)
    args_b = ["--format", "csv", "--category", "credit_card",
               "--count", "10", "--seed", "99", "--output", str(b)]
    _invoke(runner, args_a)   # seed=42
    _invoke(runner, args_b)   # seed=99
    assert a.read_text(encoding="utf-8") != b.read_text(encoding="utf-8")


# ── --evasion-rate ────────────────────────────────────────────────────────────

def test_evasion_rate_zero_no_techniques_in_csv(tmp_path):
    runner = CliRunner()
    out = tmp_path / "out.csv"
    _invoke(runner, [
        "--format", "csv", "--category", "credit_card",
        "--count", "20", "--seed", "1",
        "--evasion-rate", "0.0",
        "--output", str(out),
    ])
    with open(out, encoding="utf-8") as fh:
        rows = list(csv.DictReader(fh))
    assert all(r["technique"] == "" for r in rows)


def test_evasion_rate_one_all_have_techniques_in_csv(tmp_path):
    runner = CliRunner()
    out = tmp_path / "out.csv"
    _invoke(runner, [
        "--format", "csv", "--category", "credit_card",
        "--count", "10", "--seed", "2",
        "--evasion-rate", "1.0",
        "--output", str(out),
    ])
    with open(out, encoding="utf-8") as fh:
        rows = list(csv.DictReader(fh))
    # Credit cards have applicable generators so all should have techniques
    assert all(r["technique"] != "" for r in rows)


# ── --technique filtering ─────────────────────────────────────────────────────

def test_technique_filter_in_csv(tmp_path):
    runner = CliRunner()
    out = tmp_path / "out.csv"
    _invoke(runner, [
        "--format", "csv", "--category", "credit_card",
        "--count", "20", "--seed", "3",
        "--evasion-rate", "1.0",
        "--technique", "homoglyph_substitution",
        "--output", str(out),
    ])
    with open(out, encoding="utf-8") as fh:
        rows = list(csv.DictReader(fh))
    for r in rows:
        if r["technique"]:
            assert r["technique"] == "homoglyph_substitution"


# ── --random mode ─────────────────────────────────────────────────────────────

def test_random_mode_generates_file(tmp_path):
    runner = CliRunner()
    out = tmp_path / "out.csv"
    result = _invoke(runner, [
        "--format", "csv", "--random", "--count", "5", "--seed", "55",
        "--output", str(out),
    ])
    assert result.exit_code == 0
    assert out.exists()


# ── Multi-category ────────────────────────────────────────────────────────────

def test_multi_category_count_total(tmp_path):
    runner = CliRunner()
    out = tmp_path / "out.csv"
    _invoke(runner, [
        "--format", "csv",
        "--category", "credit_card",
        "--category", "ssn",
        "--count", "5", "--seed", "10",
        "--output", str(out),
    ])
    with open(out, encoding="utf-8") as fh:
        rows = list(csv.DictReader(fh))
    assert len(rows) == 10  # 5 per category


# ── Error cases ───────────────────────────────────────────────────────────────

def test_bad_output_path_exits_nonzero(tmp_path):
    runner = CliRunner()
    result = _invoke(runner, [
        "--format", "csv", "--category", "credit_card",
        "--count", "5", "--seed", "0",
        "--output", "/nonexistent_dir/sub/out.csv",
    ])
    assert result.exit_code != 0


def test_output_path_mentioned_in_success_output(tmp_path):
    runner = CliRunner()
    out = tmp_path / "out.csv"
    result = _invoke(runner, _base_args("csv", str(out)))
    # Path may wrap across lines in narrow terminal; collapse before checking
    assert out.name in result.output.replace("\n", "")
