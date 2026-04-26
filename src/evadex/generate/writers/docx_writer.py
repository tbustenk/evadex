"""DOCX writer for evadex generate.

Structure:
  - Title page (document name, date, disclaimer)
  - One heading-1 section per category
    - Prose paragraphs for the first two-thirds of entries
    - A table for the remaining third (shows plain / variant / technique)

Performance notes
-----------------
Table row construction bypasses python-docx's ORM (table.add_row + cell.text=)
and builds ``<w:tr>`` elements directly via lxml. This avoids the ~20× overhead
from _add_child / xmlchemy calls that python-docx incurs per cell, reducing
1 000-record DOCX generation from ~37 s to ~8 s on typical hardware.

Prose paragraphs similarly use direct lxml insertion rather than
Document.add_paragraph() to avoid per-paragraph wrapper allocation.
"""
from __future__ import annotations

import datetime
from collections import defaultdict

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from evadex.generate.generator import GeneratedEntry
from evadex.core.result import PayloadCategory

# OOXML namespace URI used for all w: elements.
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W = f"{{{_W_NS}}}"
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


_SECTION_TITLES: dict[PayloadCategory, str] = {
    PayloadCategory.CREDIT_CARD: "Credit Card Numbers",
    PayloadCategory.SSN:         "Social Security Numbers",
    PayloadCategory.SIN:         "Canadian Social Insurance Numbers",
    PayloadCategory.IBAN:        "International Bank Account Numbers",
    PayloadCategory.SWIFT_BIC:   "SWIFT / BIC Codes",
    PayloadCategory.ABA_ROUTING: "ABA Routing Numbers",
    PayloadCategory.BITCOIN:     "Bitcoin Addresses",
    PayloadCategory.ETHEREUM:    "Ethereum Addresses",
    PayloadCategory.US_PASSPORT: "US Passport Numbers",
    PayloadCategory.AU_TFN:      "Australian Tax File Numbers",
    PayloadCategory.DE_TAX_ID:   "German Tax Identification Numbers",
    PayloadCategory.FR_INSEE:    "French INSEE / NIR Numbers",
    PayloadCategory.AWS_KEY:     "AWS Access Keys",
    PayloadCategory.GITHUB_TOKEN:"GitHub Personal Access Tokens",
    PayloadCategory.STRIPE_KEY:  "Stripe API Keys",
    PayloadCategory.SLACK_TOKEN: "Slack Bot Tokens",
    PayloadCategory.JWT:         "JSON Web Tokens",
    PayloadCategory.CLASSIFICATION: "Classification Labels",
    PayloadCategory.EMAIL:       "Email Addresses",
    PayloadCategory.PHONE:       "Phone Numbers",
}


def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _make_t(parent, text: str):
    """Append a <w:t> element with preserved-space attribute when needed."""
    t = _sub(parent, "t")
    t.text = text or ""
    if text and (text[0] == " " or text[-1] == " "):
        t.set(_XML_SPACE, "preserve")
    return t


def _sub(parent, local: str):
    """Shorthand: SubElement in the w: namespace."""
    from lxml import etree
    return etree.SubElement(parent, f"{_W}{local}")


def _add_table_section(doc: Document, entries: list[GeneratedEntry], cat: PayloadCategory) -> None:
    """Write a table section using direct lxml XML construction for data rows.

    Building <w:tr> elements directly via lxml is ~15× faster than calling
    table.add_row() + cell.text= for each row, because it avoids the
    xmlchemy ORM overhead in python-docx (one lxml SubElement call vs.
    ~30 _add_child/_get_or_add_child calls per row).
    """
    from lxml import etree

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Header row — python-docx API is fine here (only 4 cells, runs once).
    hdr_cells = table.rows[0].cells
    for cell, text in zip(hdr_cells, ["#", "Value", "Technique", "Category"]):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        _shade_cell(cell, "1F4E79")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    tbl_elem = table._tbl  # lxml element for <w:tbl>

    cat_value = cat.value

    for i, e in enumerate(entries, 1):
        tr = etree.SubElement(tbl_elem, f"{_W}tr")
        col_data = (
            (str(i),                  False),
            (e.variant_value,         bool(e.technique)),
            (e.technique or "plain",  False),
            (cat_value,               False),
        )
        for text, shade in col_data:
            tc = etree.SubElement(tr, f"{_W}tc")
            if shade:
                tcPr = etree.SubElement(tc, f"{_W}tcPr")
                shd = etree.SubElement(tcPr, f"{_W}shd")
                shd.set(f"{_W}val", "clear")
                shd.set(f"{_W}color", "auto")
                shd.set(f"{_W}fill", "FFF2CC")
            p = etree.SubElement(tc, f"{_W}p")
            r = etree.SubElement(p, f"{_W}r")
            _make_t(r, text)

    doc.add_paragraph()


def _fast_add_paragraphs(doc: Document, texts: list[str]) -> None:
    """Insert multiple <w:p> elements at end of document body (before <w:sectPr>).

    python-docx's Document.add_paragraph() allocates several Python ORM
    wrappers per call — acceptable for occasional use but expensive across
    thousands of entries. Building the <w:p> elements directly via lxml and
    bulk-inserting them before <w:sectPr> is ~6× faster.
    """
    from lxml import etree

    body = doc._body._element   # lxml Element for <w:body>
    n = len(list(body))         # sectPr is always the last child
    insert_at = n - 1

    for idx, text in enumerate(texts):
        p = etree.Element(f"{_W}p")
        if text:
            r = etree.SubElement(p, f"{_W}r")
            _make_t(r, text)
        body.insert(insert_at + idx, p)


def write_docx(entries: list[GeneratedEntry], path: str) -> None:
    from evadex.generate.writers import (
        _active_template, _active_noise_level, _active_density, _active_seed,
        _active_language,
    )

    doc = Document()

    # ── Cover / title ──────────────────────────────────────────────────────
    heading = doc.add_heading("DLP Test Document", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    today = datetime.date.today().isoformat()
    sub = doc.add_paragraph(f"Generated: {today}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    disclaimer = doc.add_paragraph(
        "CONFIDENTIAL — This document contains synthetic sensitive data "
        "generated by evadex for DLP scanner testing. "
        "All values are artificial and must not be used in production."
    )
    disclaimer.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    doc.add_paragraph()

    # If a non-generic template is active, use the template system for content
    if _active_template != "generic":
        from evadex.generate.templates import apply_template
        lines = apply_template(
            _active_template, entries,
            seed=_active_seed,
            noise_level=_active_noise_level,
            density=_active_density,
            language=_active_language,
        )
        for line in lines:
            if line.startswith("===") or line.startswith("---"):
                continue
            if line.strip():
                doc.add_paragraph(line)
        doc.save(path)
        return

    by_cat: dict[PayloadCategory, list[GeneratedEntry]] = defaultdict(list)
    for e in entries:
        by_cat[e.category].append(e)

    for cat in sorted(by_cat.keys(), key=lambda c: c.value):
        cat_entries = by_cat[cat]
        title = _SECTION_TITLES.get(cat, cat.value.replace("_", " ").title())
        doc.add_heading(title, level=1)

        # Split: first ~2/3 as prose, remaining as table
        split = max(1, len(cat_entries) * 2 // 3)
        prose_entries = cat_entries[:split]
        table_entries = cat_entries[split:]

        _fast_add_paragraphs(doc, [e.embedded_text for e in prose_entries])

        if table_entries:
            _add_table_section(doc, table_entries, cat)

    doc.save(path)
