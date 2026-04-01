import io
from typing import Literal

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from fpdf import FPDF
    HAS_FPDF = True
except ImportError:
    HAS_FPDF = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME  = "application/pdf"
XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

NOISE_SENTENCES = [
    "This document contains confidential business information.",
    "Please handle with care and do not distribute.",
    "For internal use only.",
    "Authorized personnel only.",
]


class FileBuilder:
    @staticmethod
    def build(text: str, fmt: Literal["docx", "pdf", "xlsx"]) -> tuple[bytes, str]:
        """Build an in-memory document containing text. Returns (bytes, mime_type).
        Never writes to disk — uses io.BytesIO only.
        """
        if fmt == "docx":
            return FileBuilder._build_docx(text), DOCX_MIME
        elif fmt == "pdf":
            return FileBuilder._build_pdf(text), PDF_MIME
        elif fmt == "xlsx":
            return FileBuilder._build_xlsx(text), XLSX_MIME
        else:
            raise ValueError(f"Unknown format: {fmt!r}")

    @staticmethod
    def _build_docx(text: str) -> bytes:
        if not HAS_DOCX:
            raise RuntimeError("python-docx is required for DOCX generation")
        doc = Document()
        doc.add_paragraph(NOISE_SENTENCES[0])
        doc.add_paragraph(NOISE_SENTENCES[1])
        doc.add_paragraph(text)
        doc.add_paragraph(NOISE_SENTENCES[2])
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def _build_pdf(text: str) -> bytes:
        if not HAS_FPDF:
            raise RuntimeError("fpdf2 is required for PDF generation")
        pdf = FPDF()
        pdf.add_page()
        # Try DejaVu (bundled with fpdf2 >= 2.7.6) for full Unicode support.
        # Fall back to Helvetica for ASCII-only content if DejaVu is unavailable.
        try:
            pdf.set_font("DejaVu", size=12)
        except Exception:
            # DejaVu not available — encode text to latin-1 safe subset
            pdf.set_font("Helvetica", size=12)
            text = text.encode("latin-1", errors="replace").decode("latin-1")

        for sentence in NOISE_SENTENCES[:2]:
            try:
                pdf.cell(0, 10, sentence, new_x="LMARGIN", new_y="NEXT")
            except Exception:
                safe = sentence.encode("latin-1", errors="replace").decode("latin-1")
                pdf.cell(0, 10, safe, new_x="LMARGIN", new_y="NEXT")

        try:
            pdf.cell(0, 10, text, new_x="LMARGIN", new_y="NEXT")
        except Exception:
            safe = text.encode("latin-1", errors="replace").decode("latin-1")
            pdf.cell(0, 10, safe, new_x="LMARGIN", new_y="NEXT")

        pdf.cell(0, 10, NOISE_SENTENCES[2], new_x="LMARGIN", new_y="NEXT")

        buf = io.BytesIO()
        pdf.output(buf)
        return buf.getvalue()

    @staticmethod
    def _build_xlsx(text: str) -> bytes:
        if not HAS_OPENPYXL:
            raise RuntimeError("openpyxl is required for XLSX generation")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = NOISE_SENTENCES[0]
        ws['A2'] = text
        ws['A3'] = NOISE_SENTENCES[1]
        ws['B1'] = "Document ID"
        ws['B2'] = "12345"
        ws['C1'] = "Classification"
        ws['C2'] = "Confidential"
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()
